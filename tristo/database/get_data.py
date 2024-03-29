import logging

import docx
import pandas as pd
from docx import Document
from sqlalchemy import Float, case, cast, func, select
from sqlalchemy.orm import Session

from ..config import LOG_FMT
from ..status import Status
from .tables import (WVG, WVG_LAU, Data, File_Cleaned, File_Index, File_Info,
                     Param, Response)

logging.basicConfig(level=logging.INFO, format=LOG_FMT, style="{")
logger = logging.getLogger(__name__)


def get_vals(id: int, session: Session, alias=Param.param_en):
    param_query = select(case((alias!=None, alias),else_=Param.param), cast(Param.limit, Float), Param.unit).where(
        Param.id == id
    )
    param, limit, unit = session.execute(param_query).one()

    query = (
        select(
            Response.LAU,
            Data.id,
            (Data.val_num * Data.unit_factor).label("val"),
            Data.category,
        )
        .distinct()
        .join(File_Index, Data.hash2 == File_Index.hash2)
        .join(Response, Response.hash == File_Index.hash)
        .where(
            Data.param_id == id,
            Data.omitted_id == None,
            Data.val_num != None,
            Data.unit_factor != None,
        )
    )
    df = pd.read_sql(query, session.connection())

    return param, limit, unit, df


def summary(session: Session):
    results = {}

    res_name = "Supplier has n_th result:"
    total = session.execute(func.count(Response.LAU.distinct())).scalar_one() / 100
    query = (
        select(
            (Response.position + 1).label("N"),
            (cast(func.count(Response.position), Float) / total).label("Fraction"),
        )
        .where(Response.supplier.in_(["CONFIRMED", "YES"]))
        .group_by(Response.position)
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Analyzed WVG:"
    query = (
        select(func.count(WVG_LAU.wvg.distinct()))
        .join(Response)
        .where(Response.LAU != None)
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Visited Pages:"
    query = select(func.count(Response.link))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Visited unique Pages:"
    query = select(func.count(Response.link.distinct()))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Visited supplier Pages:"
    query = select(func.count(Response.link.distinct())).where(
        Response.supplier.in_([Status.YES, Status.CONFIRMED])
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Pages with downloads:"
    query = select(func.count(File_Index.hash.distinct()))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Downloaded  files:"
    query = select(func.count(File_Index.hash2))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Downloaded  Tabular-Data:"
    query = select(func.count(File_Index.hash2)).where(File_Index.ext.like(".xls%"))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Downloaded  PDF-files:"
    query = select(func.count(File_Index.hash2)).where(File_Index.ext.like(".pdf"))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Candidates for Reports:"
    query = select(func.count(File_Cleaned.hash2))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Detected Reports:"
    query = select(func.count(File_Cleaned.hash2)).where(
        File_Cleaned.status == Status.INSERTED
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Scanned Reports:"
    query = (
        select(func.count(File_Cleaned.hash2.distinct()))
        .join(File_Info, File_Info.hash2 == File_Cleaned.hash2)
        .where(File_Cleaned.status == Status.INSERTED, File_Info.status == Status.SCAN)
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Currentness of Reports:"
    query = """SELECT date 'Year',
            COUNT(date) 'Counts', 
    ROUND(100*COUNT(date)/CAST((SELECT COUNT(date) FROM file_info 
    JOIN file_cleaned ON file_info.hash2=file_cleaned.hash2 
    WHERE file_cleaned.status=='INSERTED' AND file_info.date NOT NULL) as FLOAT), 2) 'Fraction [%]'
FROM file_info 
JOIN file_cleaned ON file_info.hash2=file_cleaned.hash2 
WHERE file_cleaned.status=='INSERTED' AND file_info.date NOT NULL
GROUP BY date 
ORDER BY date DESC;"""
    # query = (
    #     select(File_Info.date, func.count(File_Info.date).label("Count"))
    #     .join(File_Cleaned, File_Cleaned.hash2 == File_Info.hash2)
    #     .where(File_Cleaned.status == Status.INSERTED, File_Info.date != None)
    #     .group_by(File_Info.date)
    #     .order_by(File_Info.date.desc())
    # )
    currentness = pd.read_sql(query, session.connection())
    results[res_name] = currentness
    currentness.to_excel("currentness of data.xlsx")

    res_name = "PDF-files from scanned Documents:"
    query = select(func.count(File_Info.hash2)).where(File_Info.status == Status.SCAN)
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "LAU per WVG:"
    sq = (
        select(func.count(WVG_LAU.LAU))
        .join(WVG, WVG.id == WVG_LAU.wvg)
        .where(WVG.supplied > 41862)
        .group_by(WVG.id)
        .order_by(WVG.supplied.desc())
        .subquery()
    )
    (count,) = tuple(sq.c)
    query = select(count, func.count(count)).group_by(count).order_by(count.desc())
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "WVG with no data:"
    query = (
        select(func.count(WVG.id.distinct()), select(func.count(WVG.id)).subquery())
        .join(WVG_LAU, WVG.id == WVG_LAU.wvg)
        .outerjoin(Response, WVG_LAU.LAU == Response.LAU)
        .where(Response.LAU == None)
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "n param per Report:"
    sq = (
        select(func.count(Data.param_id.distinct()).label("N params"))
        .group_by(Data.hash2)
        .subquery()
    )
    (count,) = tuple(sq.c)
    query = (
        select(count, func.count(count).label("N reports"))
        .group_by(count)
        .order_by(count)
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "n param per Report ungrouped:"
    query = (
        select(func.count(Data.param_id.distinct()).label("N params"))
        .group_by(Data.hash2)
        .subquery()
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Distinct Parameter:"
    query = select(func.count(Data.param_id.distinct()))
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Parameters unregulated by TrinkwV:"
    query = (
        select(func.count(Data.param_id.distinct()))
        .join(Param)
        .where(Param.origin.not_like("TrinkwV%"))
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Parameter Count:"
    query = (
        select(Param.param, func.count(Data.param_id).label("Count"))
        .join(Param)
        .where(
            Data.omitted_id == None,
            Data.val_num != None,
            Data.unit_factor != None,
            Data.val_num * Data.unit_factor <= 1.2 * cast(Param.limit, Float),
            Data.category.in_(('BG','> BG'))
        )
        .group_by(Data.param_id)
        .order_by(func.count(Data.param_id))
    )
    results[res_name] = pd.read_sql(query, session.connection())

    res_name = "Parameter Summary"
    query = """SELECT param.param 'Parameter', 
            COUNT(data.param_id) 'Value count', 
            COUNT(DISTINCT data.hash2) 'Report count',
            ROUND(100*COUNT(DISTINCT data.hash2)/(SELECT CAST(COUNT(DISTINCT hash2) AS FLOAT) FROM data WHERE data.omitted_id IS NULL AND data.val_num NOT NULL AND data.unit_factor NOT NULL), 2) 'Fraction of reports [%]', 
            IIF(param.origin NOT NULL, param.origin, '') 'Origin', 
            IIF(param.'limit' NOT NULL, param.'limit'||' '||param.unit ,'') 'Limit'
        FROM data 
        JOIN param ON data.param_id=param.id 
        WHERE data.omitted_id IS NULL AND data.val_num NOT NULL AND data.unit_factor NOT NULL 
        GROUP BY data.param_id 
        ORDER BY COUNT(data.param_id) DESC;"""
    res = pd.read_sql(query, session.connection())
    results[res_name] = res
    res.to_excel("observations_per_param.xlsx")

    return results


# https://stackoverflow.com/questions/57586400/how-to-create-bookmarks-in-a-word-document-then-create-internal-hyperlinks-to-t
def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
    start.set(docx.oxml.ns.qn("w:id"), "0")
    start.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
    end.set(docx.oxml.ns.qn("w:id"), "0")
    end.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(end)


def save_data_as_bookmarks(res: dict, fname: str):
    doc = Document()

    for key, value in res.items():
        par = doc.add_paragraph(f"{key}:")
        add_bookmark(par, f"{value:,}", key)

    doc.save(f"{fname}.docx")


def get_data_for_pub(SQL_file: str, session: Session):
    with open(f"{SQL_file}.sql", "r") as f:
        queries = f.read()

    results = {}
    for query in filter(lambda x: x.strip() != "", queries.split(";")):
        (res,) = pd.read_sql(query, session.connection()).to_dict("records")
        results.update(res)

    # calc relative values
    rel = {
        "percent_WVG_with_data": int(
            results["N_WVG_with_data"] / results["N_WVG"] * 100
        ),
        "percent_LAU_with_data": int(
            results["N_LAU_with_data"] / results["N_scraped_LAU"] * 100
        ),
        "percent_commercial": int(results["N_omitted"] / results["N_searches"] * 100),
        "percent_supplier": int(results["N_supplier"] / results["N_searches"] * 100),
        "percent_images": int(results["N_img"] / results["N_download"] * 100),
        "percent_pdf": int(results["N_pdf"] / results["N_download"] * 100),
        "percent_2022": int(results["N_2022"] / results["N_with_date"] * 100),
        "percent_2021": int(results["N_2021"] / results["N_with_date"] * 100),
        "percent_2020": int(results["N_2020"] / results["N_with_date"] * 100),
        "percent_older_2020": int(
            results["N_older_2020"] / results["N_with_date"] * 100
        ),
        'percent_file_with_limit':int(results['N_files_with_limit_col']/results['N_files_in_data'] *100),
        'percent_file_with_value_range':int(results['N_value_range']/results['N_files_in_data'] * 100),
        'percent_candidate_data':int(results['N_candidate_data']/results['N_all_data'] * 100),
        #'percent_':int(results['']/results['']*100),
    }

    results.update(rel)
    save_data_as_bookmarks(res=results, fname=SQL_file)

